import io
import pathlib
import re
from docx.oxml.ns import qn
from docx import Document
import pandas as pd
import zipfile
import json
PADRAO = r"\{\{(.+?)\}\}"

def procura_marcacoes(doc):
    """
    Procura todo o padrão de marcações no documento e retorna uma
    lista com todas as marcações encontradas.
    Exemplo: {{MARCAÇÃO1}}, {{MARCAÇÃO2}}, {{MARCAÇÃO3}}
    
    Args:
        documento (docx.Document): Documento doc recebido.
        padrao (str): str com a expressão regular para encontrar as marcações.

    Returns:
        list: lista com todas as marcações encontradas.
    """
    matches =  []
    for p in doc.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            matches.extend(re.findall(PADRAO, text))
    print("Marcações:" , matches)
    return matches

def matches_dict(doc, matches, df):
    dic = {
        "matches": {},
        "noMatches": {},
        "columns" : list(df.columns)
        }

    for match in matches:
        if match in df.columns : 
            dic["matches"][match] = match
        else:
            dic["noMatches"][match] = ""

    print(dic)
    return dic


def replace_doc(doc, dic: dict, df, file_path: pathlib.Path, bucket):
    # Create an in-memory zip archive
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for _, row in df.iterrows():
            row_dic = {f'{{{{{key}}}}}': row[key] for key in dic.keys()}
            print(row_dic)

            # Create a new document identical to the template
            new_doc = Document()
            for para in doc.paragraphs:
                get_para_data(new_doc, para)

            # Perform the substitutions
            for p in new_doc.paragraphs:
                inline = p.runs
                for i in range(len(inline)):
                    text = inline[i].text
                    for key in row_dic.keys():
                        if key in text:
                            text = text.replace(key, row_dic[key])
                            inline[i].text = text

            # Save the new document to an in-memory buffer
            new_doc_io = io.BytesIO()
            new_doc.save(new_doc_io)
            new_doc_io.seek(0)

            # Add the new document to the zip file
            new_doc_name = f"{row['Nome']}-result.docx"
            zip_file.writestr(new_doc_name, new_doc_io.read())

    # Reset the buffer position to the beginning
    zip_buffer.seek(0)

    # Define the path for the zip file in the bucket
    zip_file_path = file_path.parent.name + '/' + "results.zip"
    zip_blob = bucket.blob(str(zip_file_path))

    # Upload the zip file to Firebase Storage
    zip_blob.upload_from_string(zip_buffer.getvalue(), content_type="application/zip")



def get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        output_run.font.size = run.font.size
        output_run.font.bold = run.font.bold
        output_run.style = run.style

        

    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment